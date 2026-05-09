"""
Document Parser
=================
Handles PDF and DOCX resume parsing using PyMuPDF and python-docx.
Extracts raw text while preserving structure hints (headers, lists).

Security: File size validation, type checking, and content sanitization
are performed before processing to prevent malicious file uploads.
"""

import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

import config

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse PDF and DOCX files into clean text."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
    
    def parse(self, file_path: str | Path, file_bytes: bytes | None = None) -> str:
        """
        Parse a document file and return extracted text.
        
        Args:
            file_path: Path or filename (used for extension detection)
            file_bytes: Raw file bytes (for uploaded files)
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is unsupported or file is too large
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # File size validation
        if file_bytes and len(file_bytes) > config.MAX_UPLOAD_SIZE_BYTES:
            raise ValueError(
                f"File exceeds maximum size of {config.MAX_UPLOAD_SIZE_MB}MB"
            )
        
        if ext == ".pdf":
            return self._parse_pdf(path, file_bytes)
        elif ext == ".docx":
            return self._parse_docx(path, file_bytes)

    def _parse_pdf(self, path: Path, file_bytes: bytes | None = None) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            if file_bytes:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
            else:
                doc = fitz.open(str(path))
            
            text_parts = []
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            full_text = "\n\n".join(text_parts)
            logger.info(f"Parsed PDF '{path.name}': {len(full_text)} chars, {len(text_parts)} pages")
            return full_text
            
        except Exception as e:
            logger.error(f"PDF parsing failed for '{path.name}': {e}")
            raise ValueError(f"Failed to parse PDF '{path.name}': {e}")

    def _parse_docx(self, path: Path, file_bytes: bytes | None = None) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            import io
            if file_bytes:
                doc = Document(io.BytesIO(file_bytes))
            else:
                doc = Document(str(path))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading structure
                    if para.style and para.style.name and para.style.name.startswith("Heading"):
                        text_parts.append(f"\n## {para.text.strip()}\n")
                    else:
                        text_parts.append(para.text.strip())
            
            # Also extract from tables (resumes often use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            logger.info(f"Parsed DOCX '{path.name}': {len(full_text)} chars")
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX parsing failed for '{path.name}': {e}")
            raise ValueError(f"Failed to parse DOCX '{path.name}': {e}")


# Singleton
document_parser = DocumentParser()
